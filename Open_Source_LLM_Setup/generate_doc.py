#!/usr/bin/env python3
"""
Generate a Word document with commands and illustrative Terminal screenshots
for troubleshooting Ollama / local LLM models (Qwen, Llama).
"""
import os
from PIL import Image, ImageDraw, ImageFont

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(OUT_DIR, "images")
os.makedirs(IMG_DIR, exist_ok=True)

# ---------- Fonts ----------
def get_font(size, bold=False):
    """Try to find a monospace font, fall back to default."""
    candidates = [
        "/System/Library/Fonts/Menlo.ttc",
        "/System/Library/Fonts/Supplemental/Menlo.ttc",
        "/Library/Fonts/Menlo.ttc",
        "/System/Library/Fonts/Supplemental/Courier New.ttf",
        "/System/Library/Fonts/Courier.dfont",
    ]
    for c in candidates:
        if os.path.exists(c):
            try:
                return ImageFont.truetype(c, size)
            except Exception:
                continue
    return ImageFont.load_default()

FONT = get_font(16)
FONT_BOLD = get_font(16)
FONT_TITLE = get_font(20, bold=True)
FONT_SMALL = get_font(13)

# ---------- Colors (macOS Terminal dark theme) ----------
BG = (30, 30, 30)          # terminal background
TITLE_BG = (50, 50, 50)    # title bar
TEXT = (230, 230, 230)     # default text
GREEN = (80, 200, 120)     # prompt / success
RED = (255, 100, 100)      # errors
YELLOW = (240, 200, 80)    # warnings / headings
BLUE = (100, 180, 255)     # commands
GRAY = (160, 160, 160)     # muted

# ---------- Helpers ----------
def draw_terminal(lines, title="macbook — Terminal", width=900):
    """
    Draw a macOS Terminal window with the given lines.
    lines: list of (text, color) tuples.
    """
    char_w = 10
    line_h = 22
    pad = 20
    title_h = 40
    max_len = max((len(t) for t, _ in lines), default=40)
    content_w = max_len * char_w + pad * 2
    content_h = len(lines) * line_h + pad * 2
    width = max(width, content_w)
    height = title_h + content_h

    img = Image.new("RGB", (width, height), BG)
    d = ImageDraw.Draw(img)

    # Title bar
    d.rectangle([0, 0, width, title_h], fill=TITLE_BG)
    # Traffic lights
    for cx, col in [(18, (255, 95, 86)), (40, (255, 189, 46)), (62, (39, 201, 63))]:
        d.ellipse([cx - 6, title_h // 2 - 6, cx + 6, title_h // 2 + 6], fill=col)
    # Title text
    d.text((90, title_h // 2 - 10), title, font=FONT_SMALL, fill=GRAY)

    # Content
    y = title_h + pad
    for text, color in lines:
        d.text((pad, y), text, font=FONT, fill=color)
        y += line_h

    return img


def save(img, name):
    path = os.path.join(IMG_DIR, name)
    img.save(path)
    print(f"Saved: {path}")
    return path


# ============================================================
# CASE 1: OLLAMA IS DOWN
# ============================================================
# Image 1.1 - Check if Ollama is running
img = draw_terminal([
    ("Last login: Fri Aug  7 23:20:01 on ttys001", GRAY),
    ("macbook@MacBook ~ % ", GREEN),
    ("", TEXT),
    ("# 1. Check if Ollama is running (no output = not running)", YELLOW),
    ("macbook@MacBook ~ % ps aux | grep -i ollama | grep -v grep", BLUE),
    ("macbook@MacBook ~ % ", GREEN),
    ("", TEXT),
    ("# 2. Check if port 11434 is listening (no output = not listening)", YELLOW),
    ("macbook@MacBook ~ % lsof -iTCP:11434 -sTCP:LISTEN", BLUE),
    ("macbook@MacBook ~ % ", GREEN),
    ("", TEXT),
    ("# 3. Test the API (error = Ollama is DOWN)", YELLOW),
    ("macbook@MacBook ~ % curl -s --max-time 3 http://localhost:11434/api/tags", BLUE),
    ("curl: (7) Failed to connect to localhost port 11434: Connection refused", RED),
    ("Ollama NOT responding", RED),
], title="macbook — Terminal — Ollama Status Check")
p1 = save(img, "case1_check_status.png")

# Image 1.2 - Start Ollama
img = draw_terminal([
    ("Last login: Fri Aug  7 23:20:01 on ttys001", GRAY),
    ("macbook@MacBook ~ % ", GREEN),
    ("", TEXT),
    ("# 4. Start Ollama as a background service (auto-starts at login)", YELLOW),
    ("macbook@MacBook ~ % brew services start ollama", BLUE),
    ("==> Successfully started `ollama` (label: homebrew.mxcl.ollama)", GREEN),
    ("", TEXT),
    ("# 5. Verify the service is registered", YELLOW),
    ("macbook@MacBook ~ % brew services list | grep ollama", BLUE),
    ("ollama  started  macbook  ~/Library/LaunchAgents/homebrew.mxcl.ollama.plist", GREEN),
], title="macbook — Terminal — Start Ollama")
p2 = save(img, "case1_start_ollama.png")

# Image 1.3 - Verify Ollama is up
img = draw_terminal([
    ("Last login: Fri Aug  7 23:20:01 on ttys001", GRAY),
    ("macbook@MacBook ~ % ", GREEN),
    ("", TEXT),
    ("# 6. Confirm the API responds with the model list", YELLOW),
    ("macbook@MacBook ~ % curl -s http://localhost:11434/api/tags | jq '.models[].name'", BLUE),
    ('"qwen3:32b"', GREEN),
    ('"qwen3-coder:30b"', GREEN),
    ('"qwen2.5-coder:14b"', GREEN),
    ('"qwen2.5-coder:1.5b-base"', GREEN),
    ('"llama3.1:8b"', GREEN),
    ('"llama3.2:latest"', GREEN),
    ('"nomic-embed-text:latest"', GREEN),
    ("", TEXT),
    ("# Ollama is now UP and running ✔", GREEN),
], title="macbook — Terminal — Verify Ollama")
p3 = save(img, "case1_verify_ollama.png")


# ============================================================
# CASE 2: INSTALLED MODELS NOT WORKING
# ============================================================
# Image 2.1 - List installed models
img = draw_terminal([
    ("Last login: Fri Aug  7 23:20:01 on ttys001", GRAY),
    ("macbook@MacBook ~ % ", GREEN),
    ("", TEXT),
    ("# 1. List all installed models", YELLOW),
    ("macbook@MacBook ~ % ollama list", BLUE),
    ("NAME                       ID              SIZE      MODIFIED", GRAY),
    ("qwen3:32b                  030ee887880f    20GB      2 minutes ago", TEXT),
    ("qwen3-coder:30b            06c1097efce0    18GB      1 hour ago", TEXT),
    ("qwen2.5-coder:14b          9ec8897f747e    9.0GB     2 hours ago", TEXT),
    ("qwen2.5-coder:1.5b-base    02e0f2817a89    1.0GB     3 hours ago", TEXT),
    ("llama3.1:8b                46e0c10c039a    4.9GB     4 hours ago", TEXT),
    ("llama3.2:latest            a80c4f17acd5    2.0GB     5 hours ago", TEXT),
    ("nomic-embed-text:latest    0a109f422b47    274MB     6 hours ago", TEXT),
], title="macbook — Terminal — List Models")
p4 = save(img, "case2_list_models.png")

# Image 2.2 - Test a model with a simple prompt
img = draw_terminal([
    ("Last login: Fri Aug  7 23:20:01 on ttys001", GRAY),
    ("macbook@MacBook ~ % ", GREEN),
    ("", TEXT),
    ("# 2. Test a model with a simple prompt", YELLOW),
    ("macbook@MacBook ~ % ollama run qwen2.5-coder:14b \"Say hello\"", BLUE),
    ("Hello! How can I assist you today?", GREEN),
    ("", TEXT),
    ("# 3. Test another model (Qwen 3)", YELLOW),
    ("macbook@MacBook ~ % ollama run qwen3:32b \"What is 2+2?\"", BLUE),
    ("4", GREEN),
    ("", TEXT),
    ("# 4. Test Llama", YELLOW),
    ("macbook@MacBook ~ % ollama run llama3.1:8b \"Say hi\"", BLUE),
    ("Hi there! How can I help you today?", GREEN),
], title="macbook — Terminal — Test Models")
p5 = save(img, "case2_test_models.png")

# Image 2.3 - Test via API (curl)
img = draw_terminal([
    ("Last login: Fri Aug  7 23:20:01 on ttys001", GRAY),
    ("macbook@MacBook ~ % ", GREEN),
    ("", TEXT),
    ("# 5. Test via the API (curl) - non-streaming", YELLOW),
    ("macbook@MacBook ~ % curl -s http://localhost:11434/api/generate \\", BLUE),
    ("  -d '{\"model\":\"qwen2.5-coder:14b\",\"prompt\":\"Say hello\",\"stream\":false}' \\", BLUE),
    ("  | jq -r '.response'", BLUE),
    ("Hello! How can I assist you today?", GREEN),
    ("", TEXT),
    ("# 6. Check which model is currently loaded in memory", YELLOW),
    ("macbook@MacBook ~ % curl -s http://localhost:11434/api/ps | jq '.models[].name'", BLUE),
    ('"qwen2.5-coder:14b"', GREEN),
], title="macbook — Terminal — Test via API")
p6 = save(img, "case2_test_api.png")


# ============================================================
# BUILD THE WORD DOCUMENT
# ============================================================
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h

def add_code_block(text):
    """Add a shaded monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # light gray shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)
    return p

def add_image(path, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ---------- Title ----------
title = doc.add_heading("Ollama & Local LLM Troubleshooting Guide", 0)
sub = doc.add_paragraph()
r = sub.add_run("Commands to diagnose and fix Ollama and installed models (Qwen, Llama) on macOS")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph(
    "This guide covers two common problems: (1) Ollama is down and not working, and "
    "(2) installed models like Qwen and Llama are not working. Each section lists the "
    "commands to run, where to run them (the macOS Terminal app), and an illustrative "
    "screenshot of the expected output."
)

# ============================================================
# SECTION 1: OLLAMA IS DOWN
# ============================================================
add_heading("Case 1: Ollama is Down and Not Working", 1)

doc.add_paragraph(
    "Ollama runs as a local server on port 11434. If it is not running, any request to "
    "Qwen or Llama will fail with a connection error. Follow these steps in the "
    "macOS Terminal app (Applications → Utilities → Terminal, or Spotlight + type "
    "\u201cTerminal\u201d)."
)

add_heading("Step 1 — Check if Ollama is running", 2)
doc.add_paragraph("Run these commands to see if the Ollama process and its port are active:")
add_code_block(
    "# 1. Check if the Ollama process is running (no output = not running)\n"
    "ps aux | grep -i ollama | grep -v grep\n\n"
    "# 2. Check if port 11434 is listening (no output = not listening)\n"
    "lsof -iTCP:11434 -sTCP:LISTEN\n\n"
    "# 3. Test the API (error = Ollama is DOWN)\n"
    "curl -s --max-time 3 http://localhost:11434/api/tags"
)
add_image(p1, "Where to run: macOS Terminal — checking Ollama status (no output means it is down)", 6.0)

add_heading("Step 2 — Start Ollama as a background service", 2)
doc.add_paragraph(
    "If the checks above show Ollama is down, start it as a background service. "
    "This registers it to auto-start at login and keep running:"
)
add_code_block(
    "# Start Ollama as a background service (auto-starts at login)\n"
    "brew services start ollama\n\n"
    "# Verify the service is registered\n"
    "brew services list | grep ollama"
)
add_image(p2, "Where to run: macOS Terminal — starting Ollama with brew services", 6.0)

add_heading("Step 3 — Verify Ollama is up", 2)
doc.add_paragraph("Confirm the API responds and lists your installed models:")
add_code_block(
    "# Confirm the API responds with the model list\n"
    "curl -s http://localhost:11434/api/tags | jq '.models[].name'"
)
add_image(p3, "Where to run: macOS Terminal — verifying Ollama is up and models are listed", 6.0)

doc.add_paragraph(
    "If you see the model names listed, Ollama is now UP and running. "
    "You can now use Qwen and Llama models."
)

# ============================================================
# SECTION 2: MODELS NOT WORKING
# ============================================================
add_heading("Case 2: Installed Models (Qwen, Llama) Not Working", 1)

doc.add_paragraph(
    "If Ollama is running but a specific model fails, the model may not be installed, "
    "may be corrupted, or may be too large for your system memory. Follow these steps "
    "in the macOS Terminal app."
)

add_heading("Step 1 — List installed models", 2)
doc.add_paragraph("Check which models are actually installed:")
add_code_block(
    "# List all installed models\n"
    "ollama list"
)
add_image(p4, "Where to run: macOS Terminal — listing installed models", 6.0)

add_heading("Step 2 — Test a model with a simple prompt", 2)
doc.add_paragraph("Run a quick prompt to test each model:")
add_code_block(
    "# Test a model with a simple prompt\n"
    "ollama run qwen2.5-coder:14b \"Say hello\"\n\n"
    "# Test another model (Qwen 3)\n"
    "ollama run qwen3:32b \"What is 2+2?\"\n\n"
    "# Test Llama\n"
    "ollama run llama3.1:8b \"Say hi\""
)
add_image(p5, "Where to run: macOS Terminal — testing models with ollama run", 6.0)

add_heading("Step 3 — Test via the API (curl)", 2)
doc.add_paragraph("You can also test models programmatically via the HTTP API:")
add_code_block(
    "# Test via the API (curl) - non-streaming\n"
    "curl -s http://localhost:11434/api/generate \\\n"
    "  -d '{\"model\":\"qwen2.5-coder:14b\",\"prompt\":\"Say hello\",\"stream\":false}' \\\n"
    "  | jq -r '.response'\n\n"
    "# Check which model is currently loaded in memory\n"
    "curl -s http://localhost:11434/api/ps | jq '.models[].name'"
)
add_image(p6, "Where to run: macOS Terminal — testing models via the API", 6.0)

add_heading("Troubleshooting tips", 2)
tips = [
    "If a model is not listed by `ollama list`, it is not installed. Install it with "
    "`ollama pull <model>` (e.g. `ollama pull qwen3:32b`).",
    "If a model fails to load, it may be too large for your RAM. Check your memory with "
    "`sysctl -n hw.memsize` (bytes) and use a smaller model if needed.",
    "If Ollama is running but slow, check memory pressure with `memory_pressure`.",
    "To stop a model that is hogging memory, run `ollama stop <model>`.",
    "To see the Ollama server log, run `tail -f /opt/homebrew/var/log/ollama.log`.",
]
for t in tips:
    doc.add_paragraph(t, style="List Bullet")

# ---------- Save ----------
out_path = os.path.join(OUT_DIR, "Ollama_LLM_Troubleshooting_Guide.docx")
doc.save(out_path)
print(f"\nSaved document: {out_path}")